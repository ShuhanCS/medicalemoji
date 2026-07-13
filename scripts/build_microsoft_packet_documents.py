"""Build the external Microsoft review documents from their Markdown sources.

Usage:
    python scripts/build_microsoft_packet_documents.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "doc"
DOCUMENTS = (
    ROOT / "docs" / "strategy" / "2026-07-12-microsoft-medical-emoji-decision-brief.md",
    ROOT / "docs" / "strategy" / "2026-07-12-microsoft-medical-emoji-product-legal-clearance.md",
)

NAVY = "0B1739"
BLUE = "2563EB"
MUTED = "667085"
PALE = "EAF1FF"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        # In WordprocessingML, shading precedes cell margins and vertical
        # alignment. Insert it before those properties so the DOCX validates.
        following_tags = {
            qn("w:noWrap"),
            qn("w:tcMar"),
            qn("w:textDirection"),
            qn("w:tcFitText"),
            qn("w:vAlign"),
            qn("w:hideMark"),
        }
        for position, child in enumerate(properties):
            if child.tag in following_tags:
                properties.insert(position, shading)
                break
        else:
            properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, url: str) -> None:
    relation = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = url
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, value: str) -> None:
    value = value.replace("`", "")
    for chunk in re.split(r"(\*\*.*?\*\*)", value):
        if not chunk:
            continue
        bold = chunk.startswith("**") and chunk.endswith("**")
        text = chunk[2:-2] if bold else chunk
        run = paragraph.add_run(text)
        run.bold = bold


def set_font(run, size: float, color: str = NAVY, bold: bool | None = None) -> None:
    run.font.name = "Segoe UI"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.18)

    settings = document.settings._element
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.insert(0, zoom)
    zoom.set(qn("w:percent"), "100")

    normal = document.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    normal.font.size = Pt(9.15)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.0

    for name, size, color in (
        ("Title", 20.5, NAVY),
        ("Heading 1", 20.5, NAVY),
        ("Heading 2", 12.0, BLUE),
    ):
        style = document.styles[name]
        style.font.name = "Segoe UI Semibold"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI Semibold")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(4 if name == "Heading 2" else 0)
        style.paragraph_format.space_after = Pt(2.4)

    for style_name in ("List Bullet", "List Number"):
        list_style = document.styles[style_name]
        list_style.font.name = "Segoe UI"
        list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
        list_style.font.size = Pt(8.9)
        list_style.font.color.rgb = RGBColor.from_string(NAVY)
        list_style.paragraph_format.left_indent = Inches(0.22)
        list_style.paragraph_format.first_line_indent = Inches(-0.13)
        list_style.paragraph_format.space_after = Pt(1.5)
        list_style.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = "MEDICAL EMOJI  |  MATERIALS FOR MICROSOFT REVIEW"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, 7.4, BLUE, True)

    footer = section.footer.paragraphs[0]
    footer.text = "Prepared by Shuhan He  |  July 13, 2026  |  Independent discussion material"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_font(run, 7.2, MUTED)


def add_metadata_box(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.35)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    cell.margin_top = Inches(0.04)
    cell.margin_bottom = Inches(0.04)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        add_inline_runs(paragraph, line)
    for run in paragraph.runs:
        set_font(run, 7.8, MUTED, run.bold)


def table_cells(value: str) -> list[str]:
    return [cell.strip() for cell in value.strip().strip("|").split("|")]


def is_table_separator(value: str) -> bool:
    cells = table_cells(value)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    table.autofit = False
    total_width = 7.34
    if column_count == 3:
        widths = [1.35, 2.15, 3.84]
    else:
        widths = [total_width / column_count] * column_count

    for row_index, values in enumerate(rows):
        row = table.add_row()
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.margin_top = Inches(0.04)
            cell.margin_bottom = Inches(0.04)
            cell.margin_left = Inches(0.06)
            cell.margin_right = Inches(0.06)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F3F6FB")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, value)
            for run in paragraph.runs:
                set_font(run, 8.4, "FFFFFF" if row_index == 0 else NAVY, row_index == 0 or run.bold)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build(source: Path) -> Path:
    document = Document()
    configure_document(document)
    lines = source.read_text(encoding="utf-8").splitlines()
    metadata: list[str] = []
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            add_inline_runs(paragraph, stripped[2:])
            index += 1
            while index < len(lines) and not lines[index].strip():
                index += 1
            while index < len(lines) and lines[index].strip() and not lines[index].lstrip().startswith("#"):
                if re.match(r"^(Prepared(?: by)?|Status):", lines[index].strip()):
                    metadata.append(lines[index].strip().replace("  ", ""))
                    index += 1
                else:
                    break
            if metadata:
                add_metadata_box(document, metadata)
            continue
        if stripped.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline_runs(paragraph, stripped[3:])
            index += 1
            continue
        if stripped.startswith("### "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(1.5)
            add_inline_runs(paragraph, stripped[4:])
            for run in paragraph.runs:
                set_font(run, 9.6, NAVY, True)
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            document.add_page_break()
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            rows = [table_cells(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(table_cells(lines[index]))
                index += 1
            add_markdown_table(document, rows)
            continue
        if stripped.startswith("https://"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            add_hyperlink(paragraph, stripped)
            index += 1
            continue
        if stripped.startswith("> "):
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_shading(cell, PALE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, stripped[2:])
            for run in paragraph.runs:
                set_font(run, 8.1, NAVY, run.bold)
            index += 1
            continue
        checklist = re.match(r"^- \[ \] (.*)$", stripped)
        bullet = re.match(r"^- (.*)$", stripped)
        numbered = re.match(r"^\d+\. (.*)$", stripped)
        if checklist:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, f"☐  {checklist.group(1)}")
        elif bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet.group(1))
        elif numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, numbered.group(1))
        elif stripped.startswith("_") and len(stripped) > 20:
            paragraph = document.add_paragraph("_" * 88)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                set_font(run, 7.4, MUTED)
        else:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, stripped)
        index += 1

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                set_font(run, 9.15, NAVY, run.bold)

    properties = document.core_properties
    properties.title = lines[0].lstrip("# ")
    properties.author = "Shuhan He"
    properties.subject = "Independent materials for Microsoft review of 2026 Medical Emoji proposals"
    properties.keywords = "Medical Emoji, Microsoft, Unicode, proposal review"

    OUTPUT.mkdir(parents=True, exist_ok=True)
    destination = OUTPUT / f"{source.stem}.docx"
    document.save(destination)
    return destination


def main() -> None:
    for source in DOCUMENTS:
        destination = build(source)
        print(destination)


if __name__ == "__main__":
    main()
